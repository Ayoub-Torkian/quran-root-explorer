# -*- coding: utf-8 -*-
"""Build docx+pdf with a STATIC, page-numbered TOC (LibreOffice headless won't populate TOC fields).
Two-pass: placeholder pages fix the TOC's own length; measure real heading pages; rewrite; rebuild.
Usage: python build_toc.py <input.md> <out.docx> <out.pdf> [rtl]"""
import sys, subprocess, os, re
try:
    import fitz
except Exception:
    fitz=None
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

md, outdocx, outpdf = sys.argv[1], sys.argv[2], sys.argv[3]
rtl = len(sys.argv)>4 and sys.argv[4]=='rtl'
OUTDIR=os.path.dirname(outpdf)

raw=open(md,encoding='utf-8').read()
# strip YAML to find body; keep it for rebuild
ym=re.match(r'^---\n.*?\n---\n', raw, re.S)
yaml=ym.group(0) if ym else ''
body=raw[len(yaml):]

# collect H1/H2 headings in order (skip those inside code; none here)
heads=[]
for m in re.finditer(r'(?m)^(#{1,2})\s+(.*)$', body):
    lvl=len(m.group(1)); txt=m.group(2).strip()
    heads.append((lvl,txt))

BIDI=''.join(chr(c) for c in [0x200e,0x200f,0x202a,0x202b,0x202c,0x2066,0x2067,0x2068,0x2069])
def norm(s):
    s=s.translate({ord(c):None for c in BIDI})
    s=re.sub(r'[*_`]','',s)            # strip md emphasis
    s=re.sub(r'\s+',' ',s).strip()
    return s

def matchkey(s):
    s=s.translate({ord(c):None for c in BIDI}).lower()
    s=re.sub(r"[^0-9a-z\u0600-\u06ff ]+"," ",s)   # keep latin/arabic/digits/space
    return re.sub(r"\s+"," ",s).strip()

def build_docx(toc_pages):
    # toc_pages: list aligned with heads giving page-number string (or placeholder)
    ref="/tmp/ref_%s.docx"%("rtl" if rtl else "ltr")
    subprocess.run(['pandoc','--print-default-data-file','reference.docx'],stdout=open(ref,'wb'),check=True)
    d=Document(ref)
    for sname in ['Normal','Body Text','First Paragraph','Compact','Footer','Header']:
        try:
            pf=d.styles[sname].paragraph_format; pf.line_spacing=1.0; pf.space_after=Pt(6); pf.space_before=Pt(0)
        except KeyError: pass
    try:
        n=d.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11)
    except KeyError: pass
    # footer page number centered
    sec=d.sections[0]; foot=sec.footer; foot.is_linked_to_previous=False
    p=foot.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run()
    for t,val in [('begin',None),('instr','PAGE'),('end',None)]:
        if t=='instr':
            e=OxmlElement('w:instrText'); e.set(qn('xml:space'),'preserve'); e.text=' PAGE '
        else:
            e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),t)
        r._r.append(e)
    d.save(ref)
    cmd=['pandoc', md, '-o', outdocx, '--reference-doc', ref]
    cmd += (['-V','dir=rtl','-V','lang=fa'] if rtl else ['-V','lang=en'])
    subprocess.run(cmd, capture_output=True, text=True)
    # now inject a static TOC right after the title block (before first heading)
    d=Document(outdocx)
    body_el=d.element.body
    # find index of first paragraph that is a heading
    title_txt = "فهرست مطالب" if rtl else "Table of Contents"
    # build TOC paragraphs and insert at top (after any leading title paragraphs from pandoc metadata)
    # We insert before the FIRST Heading-styled paragraph.
    first_head_idx=None
    paras=d.paragraphs
    for i,pp in enumerate(paras):
        if pp.style.name.startswith('Heading') or norm(pp.text)==norm(heads[0][1]):
            first_head_idx=i; break
    anchor_par = paras[first_head_idx] if first_head_idx is not None else paras[-1]
    anchor = anchor_par._p
    PARENT = anchor_par._parent
    def mkpar(text,page,lvl):
        p=OxmlElement('w:p'); anchor.addprevious(p)
        para=d.paragraphs[0]  # not used
        from docx.text.paragraph import Paragraph
        P=Paragraph(p,PARENT)
        pf=P.paragraph_format; pf.line_spacing=1.0; pf.space_after=Pt(2)
        # tab stop with dot leader near right margin
        try:
            P.paragraph_format.tab_stops.add_tab_stop(Pt(468), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        except Exception: pass
        if rtl:
            bd=OxmlElement('w:bidi'); P._p.get_or_add_pPr().append(bd)
        run=P.add_run(("    " if lvl==2 else "")+text); run.font.size=Pt(12)
        if lvl==1: run.font.bold=True
        run.font.name='Calibri'
        tab=P.add_run('\t'+str(page)); tab.font.size=Pt(12); tab.font.name='Calibri'
        return P
    # heading for the TOC
    th=OxmlElement('w:p'); anchor.addprevious(th)
    from docx.text.paragraph import Paragraph
    TH=Paragraph(th,PARENT); rr=TH.add_run(title_txt); rr.font.size=Pt(16); rr.font.bold=True; rr.font.name='Calibri'
    if rtl:
        bd=OxmlElement('w:bidi'); TH._p.get_or_add_pPr().append(bd)
    for (lvl,txt),pg in zip(heads,toc_pages):
        mkpar(norm(txt),pg,lvl)
    # page break after TOC
    pb=OxmlElement('w:p'); anchor.addprevious(pb)
    PB=Paragraph(pb,PARENT); brk=PB.add_run(); be=OxmlElement('w:br'); be.set(qn('w:type'),'page'); brk._r.append(be)
    d.save(outdocx)

def to_pdf():
    subprocess.run(['soffice','--headless','--convert-to','pdf','--outdir',OUTDIR,outdocx],
                   capture_output=True,timeout=300)

def _pdftext_pages():
    txt=subprocess.run(['pdftotext','-layout',outpdf,'-'],capture_output=True,text=True).stdout
    return [matchkey(pg) for pg in txt.split('\f')]

def heading_pages():
    full=_pdftext_pages(); npages=len(full)
    # 1) outline anchors (accurate, monotonic, but may miss some headings)
    anchors={}  # head-index -> page
    if fitz is not None:
        try:
            otoc=[(l,norm(tt),pg) for l,tt,pg in fitz.open(outpdf).get_toc() if l<=2]
        except Exception:
            otoc=[]
        oi=0
        for hi,(lvl,h) in enumerate(heads):
            nh=norm(h)
            if oi<len(otoc):
                ot=otoc[oi][1]
                if nh==ot or nh.startswith(ot[:18]) or ot.startswith(nh[:18]):
                    anchors[hi]=otoc[oi][2]; oi+=1
    # 2) fill gaps via windowed token-subset text search (bounded -> no far false positives)
    res=[0]*len(heads)
    for hi in range(len(heads)):
        if hi in anchors: res[hi]=anchors[hi]; continue
        # window between nearest known anchors
        lo=max((anchors[k] for k in anchors if k<hi), default=1)
        hi_pg=min((anchors[k] for k in anchors if k>hi), default=npages)
        toks=[w for w in matchkey(heads[hi][1]).split() if re.search(r'[a-z\u0600-\u06ff]',w) and len(w)>=2]
        found=0
        if toks:
            for pi in range(lo-1, min(hi_pg, npages)):
                pset=set(full[pi].split())
                if all(w in pset for w in toks): found=pi+1; break
        res[hi]=found or lo
    # 3) enforce monotonic non-decreasing (document order)
    for i in range(1,len(res)):
        if res[i]<res[i-1]: res[i]=res[i-1]
    return res

# PASS 1: placeholder pages (fixes TOC length)
build_docx(['00']*len(heads)); to_pdf()
pages=heading_pages()
# PASS 2: real pages
build_docx([str(p) for p in pages]); to_pdf()
pages2=heading_pages()
mism=sum(1 for a,b in zip(pages,pages2) if a!=b)
print("TOC built. entries:",len(heads)," page-stability mismatches pass1->pass2:",mism)
print("sample:", list(zip([h[1][:24] for h in heads[:6]], pages2[:6])))
print("exists pdf:", os.path.exists(outpdf), os.path.getsize(outpdf))
