# root_dossier.py — per-root trilingual (EN / FA / AR) research dossier + DOCX export.
# Fully generic: works for ANY input root. Every figure is computed from the corpus;
# lexical senses/glosses are curated (concept_senses.json); no per-root hardcoding.
import os, io, json, datetime, re
from collections import Counter
import analysis as _A

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
_DIA = re.compile(r"[ؐ-ًؚ-ٰٟۖ-ۭـ]")
_TRANSLIT = {"ا":"ʾ","أ":"ʾ","إ":"ʾ","آ":"ʾ","ء":"ʾ","ب":"b","ت":"t","ث":"th","ج":"j","ح":"ḥ",
    "خ":"kh","د":"d","ذ":"dh","ر":"r","ز":"z","س":"s","ش":"sh","ص":"ṣ","ض":"ḍ","ط":"ṭ","ظ":"ẓ",
    "ع":"ʿ","غ":"gh","ف":"f","ق":"q","ك":"k","ک":"k","ل":"l","م":"m","ن":"n","ه":"h","و":"w","ی":"y","ي":"y","ة":"a"}

def _ar(s):
    return s.replace("ک","ك").replace("ی","ي").replace("ۀ","ة") if isinstance(s,str) else s
def _bare(t):
    return _A.normalize_letters(_DIA.sub("", t or ""))
def _translit(root):
    return "-".join(_TRANSLIT.get(ch,"") for ch in _ar(root) if ch.strip())
def _spaced(root):
    return " · ".join(list(_ar(root)))

def _load_senses():
    try:
        return json.load(open(os.path.join(_HERE,"concept_senses.json"), encoding="utf-8"))
    except Exception:
        return {}

def compute(corpus, root, normalize=True):
    df = corpus.df; nz = _A.normalize_letters
    NUZ = "ترتیب نزول"; has_nuz = NUZ in df.columns
    snz = df.drop_duplicates(_A.COL_SURAH).set_index(_A.COL_SURAH)[NUZ].to_dict() if has_nuz else {}
    snames = df.drop_duplicates(_A.COL_SURAH).set_index(_A.COL_SURAH)[_A.COL_SURAH_NAME].to_dict()
    aps = df.groupby(_A.COL_SURAH).size().to_dict()
    def per(s): return "Meccan" if int(snz.get(int(s),200)) <= 86 else "Medinan"
    senses_all = _load_senses()
    S = next((senses_all[k] for k in senses_all if nz(k)==nz(root)), {})
    def gloss_of(nzr):
        node = next((senses_all[k] for k in senses_all if nz(k)==nzr), None)
        if node and node.get("senses"):
            return node["senses"][0].get("sense","")
        return ""
    idxs = _A.search_root(corpus, root, normalize)
    ayahs=set(); freq=0; surahs=set(); formc=Counter(); by_surah=Counter()
    mec=0; med=0; nuzc=Counter()
    Lc=Counter(); Rc=Counter(); bg=Counter(); tg=Counter(); partners=Counter()
    positions=[]; lens=[]; occ_positions=[]
    for i in idxs:
        r=df.iloc[i]; rt=corpus.root_tokens[i]; stk=corpus.surface_tokens[i]
        sid=int(r[_A.COL_SURAH]); aid=int(r[_A.COL_AYAH])
        ayahs.add((sid,aid)); surahs.add(sid); by_surah[sid]+=1
        if per(sid)=="Meccan": mec+=1
        else: med+=1
        if has_nuz: nuzc[int(snz[sid])]+=1
        for t in set(nz(x) for x in rt):
            if t!=root: partners[t]+=1
        hitpos=[k for k,t in enumerate(rt) if nz(t)==root]
        if hitpos: occ_positions.append((sid,aid,i,len(hitpos)))
        for j in hitpos:
            freq+=1; f=stk[j] if j<len(stk) else ""
            formc[f]+=1; positions.append(j/max(len(rt)-1,1)); lens.append(len(rt))
            L=stk[j-1] if j-1>=0 else None; Rn=stk[j+1] if j+1<len(stk) else None
            if L is not None: Lc[L]+=1
            if Rn is not None: Rc[Rn]+=1
            if L is not None: bg[L+" "+f]+=1
            if Rn is not None: bg[f+" "+Rn]+=1
            if L is not None and Rn is not None: tg[L+" "+f+" "+Rn]+=1
    nA=len(ayahs); nS=len(surahs)
    if nA==0:
        return {"root":_ar(root), "empty":True}
    TOK=sum(len(t) for t in corpus.root_tokens)
    conc=_A.root_concentration(corpus, root, normalize)
    allf=Counter()
    for toks in corpus.root_tokens:
        for t in toks: allf[nz(t)]+=1
    ranked=sorted(allf.values(), reverse=True); myf=allf[nz(root)]
    rank=sum(1 for v in ranked if v>myf)+1; nroots=len(ranked)
    pct=round(100*(1-rank/nroots),1)
    dens=[(sid,round(1000*h/int(aps[sid]),1),h,int(aps[sid]))
          for sid,h in by_surah.items() if aps.get(sid) and h>=3 and int(aps[sid])>=10]
    dens=sorted(dens,key=lambda x:-x[1])[:6]
    base=Counter(per(s) for s in df[_A.COL_SURAH]); bmec=base["Meccan"]; bmed=base["Medinan"]
    dmec=round(1000*mec/bmec,1) if bmec else 0.0
    dmed=round(1000*med/bmed,1) if bmed else 0.0
    pc=[sum(v for k,v in nuzc.items() if a<=k<=b) for a,b in [(1,20),(21,50),(51,86),(87,114)]]
    lean = "Medinan" if dmed>dmec else "Meccan"
    ratio = round(max(dmec,dmed)/max(min(dmec,dmed),0.1),1)
    # concordance: up to 4 occurrences spread across the sorted list, bold matched form
    occ_sorted=sorted(occ_positions)
    picks=[]
    if occ_sorted:
        n=len(occ_sorted)
        for frac in (0.0,0.34,0.67,0.95):
            picks.append(occ_sorted[min(int(frac*n), n-1)])
        seen=set(); uniq=[]
        for p in picks:
            if (p[0],p[1]) not in seen:
                seen.add((p[0],p[1])); uniq.append(p)
        picks=uniq[:4]
    forms_set=set(_bare(f) for f in formc)
    conc_rows=[]
    for sid,aid,i,_h in picks:
        txt=df.iloc[i][_A.COL_DIACRITIZED]
        runs=[]
        for tok in str(txt).split():
            bt=_bare(tok); hit=any(bf and bf in bt for bf in forms_set)
            runs.append((_ar(tok), hit))
        conc_rows.append({"ref":f"S{sid}:{aid}", "runs":runs})
    top_phrases=[( _ar(p), n) for p,n in sorted(bg.most_common(8)+tg.most_common(4), key=lambda x:-x[1])[:10]]
    return dict(empty=False, root=_ar(root), spaced=_spaced(root), translit=_translit(root),
        senses=S.get("senses",[]), nA=nA, nS=nS, freq=freq, nforms=len(formc),
        rank=rank, nroots=nroots, pct=pct, rate_ay=round(1000*nA/corpus.n_ayahs,1),
        rate_rt=round(1000*freq/TOK,2), gini=conc["gini"], top3=conc["top3_share"],
        meanps=round(nA/nS,1), rep=round(freq/nA,2),
        raw_top=[(_ar(snames[s]), s, h) for s,h in by_surah.most_common(6)],
        dens=[(_ar(snames[s]), s, d, h, sz) for s,d,h,sz in dens],
        forms=[(_ar(f), n, round(100*n/freq,1)) for f,n in formc.most_common(6)],
        mec=mec, med=med, dmec=dmec, dmed=dmed, bmec=bmec, bmed=bmed,
        pc=pc, lean=lean, ratio=ratio,
        left=[(_ar(w),n) for w,n in Lc.most_common(6)],
        right=[(_ar(w),n) for w,n in Rc.most_common(6)],
        phrases=top_phrases,
        partners=[(_ar(k), v, gloss_of(k)) for k,v in partners.most_common(10)],
        pos=round(sum(positions)/len(positions),2), avglen=round(sum(lens)/len(lens),1),
        concordance=conc_rows)


# ─────────────────────────────────────────────────────────────────────────
# Trilingual content (EN comprehensive; FA/AR faithful parallel) + DOCX build
# ─────────────────────────────────────────────────────────────────────────
_LEAN = {"en":{"Meccan":"Meccan","Medinan":"Medinan"},
         "fa":{"Meccan":"مکی","Medinan":"مدنی"},
         "ar":{"Meccan":"مكي","Medinan":"مدني"}}
_PHASE = {"en":["early Meccan","mid Meccan","late Meccan","Medinan"],
          "fa":["مکی نخستین","مکی میانی","مکی پایانی","مدنی"],
          "ar":["مكي مبكر","مكي أوسط","مكي متأخر","مدني"]}
_L = {
 "en":{"title":"Root Dossier","scope":"Every figure is computed from the corpus (6,236 āyāt). Lexical senses are curated; verse text is quoted, not translated. Interpretation of meaning is left to the reader.",
   "s_glance":"At a glance","s_syn":"Executive synthesis","s_senses":"Lexical senses (curated)","s_freq":"Frequency & rarity","s_dist":"Distribution across the muṣḥaf","s_revel":"Revelation timing (nuzūl)","s_forms":"Surface forms","s_colloc":"Collocations (immediate adjacency)","s_part":"Conceptual neighbourhood (partner roots)","s_shape":"Position & shape","s_conc":"Concordance — representative verses","s_method":"Method & definitions",
   "h_sense":["Sense","Occurrences","Leading forms"],"h_metric":["Metric","Value"],
   "h_dens":["Sūra","Density /1k","Hits","Sūra āyāt"],"h_rev":["Period","Āyah-hits","Density /1k"],
   "h_form":["Form","Occurrences","Share"],"h_phr":["Phrase","Count"],"h_part":["Root","Co-occurring āyāt","Gloss"],
   "m":["Āyāt (doc freq)","Total occurrences","Sūras","Surface forms","Rank / percentile","Rate /1k āyāt","Rate /1k roots","Gini / Top-3 share","Mean hits / sūra","Repetition"],
   "meccan":"Meccan (rev ≤86)","medinan":"Medinan (rev 87–114)","prec":"Preceding words","foll":"Following words",
   "pos":"Mean relative position in the āyah","posl0":"0 = start, 1 = end","alen":"Mean length of host āyah","rtoks":"root-tokens",
   "reach":"Reach","rar":"Rarity","lean_l":"Revelation lean","sig":"Signature phrase","busiest":"Busiest sūras (raw hits)","phasec":"Phase curve (āyah-hits)","denserw":"denser in","method":[
     "Corpus: 6,236 āyāt, 1,701 roots; each verse stored diacritized → segmented → root tokens.",
     "Density (size-true): occurrences per 1,000 container units (per sūra / per revelation period).",
     "Reliability floor: a sūra's density counts only at ≥3 hits and ≥10 āyāt.",
     "Revelation order: Egyptian-standard nuzūl; Meccan = revealed 1–86, Medinan = 87–114.",
     "Collocation: immediate neighbours on segmented tokens; clitics appear as separate words.",
     "Senses/glosses: curated (concept_senses.json). Verse text is quoted; renderings are the reader's."]},
 "fa":{"title":"کارنامهٔ ریشه","scope":"همهٔ اعداد از پیکرهٔ متن (۶۲۳۶ آیه) محاسبه شده‌اند. معانی ریشه برگزیده‌اند؛ متن آیات نقل شده و ترجمه نشده است. تفسیرِ معنا بر عهدهٔ خواننده است.",
   "s_glance":"نگاه کلی","s_syn":"خلاصهٔ تحلیلی","s_senses":"معانی ریشه (برگزیده)","s_freq":"بسامد و کمیابی","s_dist":"پراکندگی در مصحف","s_revel":"زمان نزول (ترتیب نزول)","s_forms":"صورت‌های ظاهری","s_colloc":"هم‌نشینی (مجاورتِ بی‌واسطه)","s_part":"همسایگی مفهومی (ریشه‌های هم‌آیند)","s_shape":"جایگاه و شکل","s_conc":"شواهد آیات","s_method":"روش و تعاریف",
   "h_sense":["معنا","بسامد","صورت‌های اصلی"],"h_metric":["سنجه","مقدار"],
   "h_dens":["سوره","چگالی/هزار","اصابت","آیات سوره"],"h_rev":["دوره","اصابت آیه","چگالی/هزار"],
   "h_form":["صورت","بسامد","سهم"],"h_phr":["عبارت","شمار"],"h_part":["ریشه","آیات مشترک","معنا"],
   "m":["آیات (بسامد سندی)","کل رخداد","سوره‌ها","صورت‌های ظاهری","رتبه/صدک","نرخ/هزار آیه","نرخ/هزار ریشه","جینی/سهمِ سه‌تای برتر","میانگین در هر سوره","تکرار"],
   "meccan":"مکی (نزول ≤۸۶)","medinan":"مدنی (نزول ۸۷–۱۱۴)","prec":"واژه‌های پیشین","foll":"واژه‌های پسین",
   "pos":"میانگین جایگاه نسبی در آیه","posl0":"۰ = آغاز، ۱ = پایان","alen":"میانگین طول آیهٔ میزبان","rtoks":"واحد ریشه",
   "reach":"گستره","rar":"کمیابی","lean_l":"گرایش نزولی","sig":"عبارت شاخص","busiest":"پرتکرارترین سوره‌ها (اصابت خام)","phasec":"منحنی مراحل (اصابت آیه)","denserw":"چگال‌تر در","method":[
     "پیکره: ۶۲۳۶ آیه، ۱۷۰۱ ریشه؛ هر آیه به‌صورت باحرکت ← توکن‌شده ← ریشه ذخیره شده است.",
     "چگالی (اندازه‌راست): رخداد در هر ۱۰۰۰ واحدِ ظرف (در هر سوره / در هر دورهٔ نزول).",
     "کفِ اطمینان: چگالی سوره تنها با ≥۳ اصابت و ≥۱۰ آیه معتبر شمرده می‌شود.",
     "ترتیب نزول: نزولِ استاندارد مصری؛ مکی = نزول ۱–۸۶، مدنی = ۸۷–۱۱۴.",
     "هم‌نشینی: همسایه‌های بی‌واسطه بر پایهٔ توکن‌های تجزیه‌شده؛ پی‌بست‌ها واژهٔ جداگانه‌اند.",
     "معانی: برگزیده (concept_senses.json). متن آیات نقل شده است؛ برگردانِ معنا با خواننده است."]},
 "ar":{"title":"ملف الجذر","scope":"كل الأرقام محسوبة من المدوّنة (٦٢٣٦ آية). المعاني منتقاة؛ نصّ الآيات منقول لا مترجَم. تفسير المعنى متروك للقارئ.",
   "s_glance":"نظرة عامة","s_syn":"الخلاصة التحليلية","s_senses":"معاني الجذر (منتقاة)","s_freq":"التواتر والندرة","s_dist":"التوزّع في المصحف","s_revel":"زمن النزول (ترتيب النزول)","s_forms":"الصور الصرفية","s_colloc":"التلازم اللفظي (المجاورة المباشرة)","s_part":"الجوار المفاهيمي (الجذور المقارِنة)","s_shape":"الموضع والشكل","s_conc":"الشواهد من الآيات","s_method":"المنهج والتعريفات",
   "h_sense":["المعنى","التواتر","أبرز الصور"],"h_metric":["المقياس","القيمة"],
   "h_dens":["السورة","الكثافة/ألف","الإصابات","آيات السورة"],"h_rev":["الفترة","إصابات الآيات","الكثافة/ألف"],
   "h_form":["الصورة","التواتر","النسبة"],"h_phr":["العبارة","العدد"],"h_part":["الجذر","الآيات المشتركة","المعنى"],
   "m":["الآيات (تواتر مستندي)","مجموع الورود","السور","الصور الصرفية","الرتبة/المئين","المعدّل/ألف آية","المعدّل/ألف جذر","جيني/نصيب الثلاث","متوسط لكل سورة","التكرار"],
   "meccan":"مكي (نزول ≤٨٦)","medinan":"مدني (نزول ٨٧–١١٤)","prec":"الكلمات السابقة","foll":"الكلمات اللاحقة",
   "pos":"متوسط الموضع النسبي في الآية","posl0":"٠ = البداية، ١ = النهاية","alen":"متوسط طول الآية المضيفة","rtoks":"وحدة جذرية",
   "reach":"المدى","rar":"الندرة","lean_l":"الميل النزولي","sig":"العبارة المميّزة","busiest":"أكثر السور (إصابات خام)","phasec":"منحنى المراحل (إصابات الآيات)","denserw":"أكثف في","method":[
     "المدوّنة: ٦٢٣٦ آية، ١٧٠١ جذرًا؛ كل آية مخزّنة مشكّلة ← مقطّعة ← جذورًا.",
     "الكثافة (المصحَّحة بالحجم): ورود لكل ١٠٠٠ وحدة (لكل سورة / لكل فترة نزول).",
     "عتبة الموثوقية: كثافة السورة تُعتمد فقط عند ≥٣ إصابات و≥١٠ آيات.",
     "ترتيب النزول: النزول المصري القياسي؛ مكي = نزول ١–٨٦، مدني = ٨٧–١١٤.",
     "التلازم: الجيران المباشرون على التوكنات المقطّعة؛ السوابق واللواحق كلمات مستقلة.",
     "المعاني: منتقاة (concept_senses.json). نصّ الآيات منقول؛ ونقل المعنى متروك للقارئ."]},
}

def _fmt_senses(D):
    out=[]
    for s in D["senses"]:
        if s.get("status")=="candidate" and s.get("occ",0)<=1: continue
        forms=", ".join(f"{_ar(f)} ({n})" for f,n in s.get("forms",[])[:4])
        out.append((_ar(s.get("sense","")), s.get("occ",0), forms))
    return out

def _blocks(lang, D):
    T=_L[lang]; lean=_LEAN[lang][D["lean"]]; B=[]
    B.append(("h1", f"{T['title']} — {D['spaced']}  ({D['translit']})"))
    B.append(("scope", T["scope"]))
    # at a glance
    B.append(("h2", T["s_glance"]))
    sc=_fmt_senses(D)
    senses_inline=" · ".join(f"{x[0]} ({x[1]})" for x in sc) if sc else "—"
    B.append(("bul",[
        f"{T['reach']}: {D['nA']} · {D['nS']}/114 · {D['freq']} · {D['nforms']}",
        f"{T['rar']}: #{D['rank']} / {D['nroots']} — {D['pct']}",
        f"{T['s_senses']}: {senses_inline}",
        f"{T['lean_l']}: {lean} (~{D['ratio']}×)",
        f"{T['sig']}: {D['phrases'][0][0]} ({D['phrases'][0][1]})" if D['phrases'] else "",
    ]))
    # synthesis
    B.append(("h2", T["s_syn"]))
    p1s = (f" It carries {len(sc)} curated sense(s): {senses_inline}." if lang=="en" else
           f" این ریشه دارای {len(sc)} معنای برگزیده است: {senses_inline}." if lang=="fa" else
           f" وللجذر {len(sc)} من المعاني المنتقاة: {senses_inline}.") if sc else ""
    gword = ({"en":("highly uneven" if D['gini']>=0.6 else "moderately uneven" if D['gini']>=0.4 else "relatively even"),
              "fa":("بسیار ناهموار" if D['gini']>=0.6 else "نسبتاً ناهموار" if D['gini']>=0.4 else "نسبتاً هموار"),
              "ar":("متفاوت جدًّا" if D['gini']>=0.6 else "متفاوت نسبيًّا" if D['gini']>=0.4 else "متساوٍ نسبيًّا")})[lang]
    mp=round(100*D['bmec']/6236); dp=round(100*D['bmed']/6236)
    phr3=", ".join(f"{p} ({n})" for p,n in D['phrases'][:3])
    part3="، ".join(p[0] for p in D['partners'][:5])
    home=D['dens'][0] if D['dens'] else None
    homes = f"{home[0]} (S{home[1]}) {home[2]}/1k" if home else "—"
    if lang=="en":
        B.append(("p",[(f"The root {D['spaced']} occurs {D['freq']} times across {D['nA']} āyāt in {D['nS']} of 114 sūras — rank #{D['rank']} of {D['nroots']} roots ({D['pct']}th percentile). Its spread is {gword} (Gini {D['gini']}; the three busiest sūras hold {D['top3']}% of all hits).{p1s}",False)]))
        B.append(("p",[(f"By revelation timing it is {lean}-leaning: {D['dmed']} vs {D['dmec']} occurrences per 1,000 āyāt (Medinan vs Meccan), about {D['ratio']}× {T['denserw']} {lean} text once the {mp}%/{dp}% Meccan/Medinan split is accounted for; phase counts run {' → '.join(str(x) for x in D['pc'])}. Its most frequent adjacent phrases are {phr3}. The size-true density home is {homes}, and its top co-occurring roots are {part3}.",False)]))
    elif lang=="fa":
        B.append(("p",[(f"ریشهٔ {D['spaced']} در {D['freq']} مورد و در {D['nA']} آیه از {D['nS']} سوره (از ۱۱۴) به‌کار رفته است — رتبهٔ #{D['rank']} از {D['nroots']} ریشه (صدکِ {D['pct']}). پراکندگی آن {gword} است (جینی {D['gini']}؛ سه سورهٔ پرتکرار {D['top3']}٪ از کل اصابت‌ها را دارند).{p1s}",False)]))
        B.append(("p",[(f"از نظر زمان نزول، این ریشه {lean}‌گراست: {D['dmed']} در برابر {D['dmec']} رخداد در هر ۱۰۰۰ آیه (مدنی در برابر مکی)، یعنی حدود {D['ratio']}× {T['denserw']} متن {lean}، پس از لحاظ‌کردن نسبتِ {mp}٪/{dp}٪ مکی/مدنی؛ شمارشِ مراحل: {' → '.join(str(x) for x in D['pc'])}. پربسامدترین عبارت‌های مجاور: {phr3}. خانهٔ چگالیِ اندازه‌راست: {homes}؛ و پربسامدترین ریشه‌های هم‌آیند: {part3}.",False)]))
    else:
        B.append(("p",[(f"يرد الجذر {D['spaced']} {D['freq']} مرة في {D['nA']} آية ضمن {D['nS']} من ١١٤ سورة — الرتبة #{D['rank']} من {D['nroots']} جذرًا (المئين {D['pct']}). توزّعه {gword} (جيني {D['gini']}؛ أكثر ثلاث سور تحوز {D['top3']}٪ من الإصابات).{p1s}",False)]))
        B.append(("p",[(f"ومن حيث زمن النزول يميل إلى {lean}: {D['dmed']} مقابل {D['dmec']} ورودًا لكل ١٠٠٠ آية (مدني مقابل مكي)، أي أكثف بنحو {D['ratio']}× في النصّ {lean} بعد مراعاة نسبة {mp}٪/{dp}٪ المكي/المدني؛ وأعداد المراحل: {' ← '.join(str(x) for x in D['pc'])}. وأكثر العبارات المجاورة: {phr3}. وموطن الكثافة المصحَّحة: {homes}؛ وأكثر الجذور اقترانًا: {part3}.",False)]))
    # senses table
    if sc:
        B.append(("h2", T["s_senses"]))
        B.append(("table", T["h_sense"], [[a,str(b),c] for a,b,c in sc]))
    # freq
    B.append(("h2", T["s_freq"]))
    mvals=[f"{D['nA']}", f"{D['freq']}", f"{D['nS']}/114", f"{D['nforms']}", f"#{D['rank']} · {D['pct']}",
           f"{D['rate_ay']}", f"{D['rate_rt']}", f"{D['gini']} · {D['top3']}%", f"{D['meanps']}", f"{D['rep']}"]
    B.append(("table", T["h_metric"], [[T["m"][i], mvals[i]] for i in range(len(mvals))]))
    # distribution
    B.append(("h2", T["s_dist"]))
    B.append(("p",[(T["busiest"]+": "+", ".join(f"{n} (S{s}) {h}" for n,s,h in D["raw_top"]),False)]))
    if D["dens"]:
        B.append(("table", T["h_dens"], [[f"{n} (S{s})", str(d), str(h), str(sz)] for n,s,d,h,sz in D["dens"]]))
    # revelation
    B.append(("h2", T["s_revel"]))
    B.append(("table", T["h_rev"], [[T["meccan"], str(D["mec"]), str(D["dmec"])],[T["medinan"], str(D["med"]), str(D["dmed"])]]))
    B.append(("p",[(T["phasec"]+": "+" → ".join(f"{_PHASE[lang][i]} {D['pc'][i]}" for i in range(4)),False)]))
    # forms
    B.append(("h2", T["s_forms"]))
    B.append(("table", T["h_form"], [[f, str(n), f"{sh}%"] for f,n,sh in D["forms"]]))
    # collocations
    B.append(("h2", T["s_colloc"]))
    B.append(("p",[(T["prec"]+": "+", ".join(f"{w} ({n})" for w,n in D["left"]),False)]))
    B.append(("p",[(T["foll"]+": "+", ".join(f"{w} ({n})" for w,n in D["right"]),False)]))
    B.append(("table", T["h_phr"], [[p, str(n)] for p,n in D["phrases"]]))
    # partners
    B.append(("h2", T["s_part"]))
    B.append(("table", T["h_part"], [[k, str(v), g] for k,v,g in D["partners"]]))
    # shape
    B.append(("h2", T["s_shape"]))
    B.append(("bul",[f"{T['pos']}: {D['pos']} ({T['posl0']})", f"{T['alen']}: {D['avglen']} {T['rtoks']}"]))
    # concordance
    B.append(("h2", T["s_conc"]))
    for row in D["concordance"]:
        B.append(("verse", row["ref"], row["runs"]))
    # method
    B.append(("h2", T["s_method"]))
    B.append(("bul", T["method"]))
    return B

# ---- DOCX rendering helpers ----
_INK=RGBColor(0x1A,0x1A,0x1A); _TEAL=RGBColor(0x0F,0x5C,0x6B); _BLUE=RGBColor(0x07,0x55,0x7A); _MUT=RGBColor(0x5A,0x6B,0x72)
def _set_rtl_p(p):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:bidi'); pPr.append(b)
def _run(p, text, rtl, bold=False, size=11, color=None, italic=False, cs="Arial"):
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color is not None: r.font.color.rgb=color
    rpr=r._r.get_or_add_rPr()
    rf=OxmlElement('w:rFonts'); rf.set(qn('w:cs'), cs); rf.set(qn('w:ascii'),"Calibri"); rf.set(qn('w:hAnsi'),"Calibri"); rpr.append(rf)
    if rtl:
        rtl_el=OxmlElement('w:rtl'); rtl_el.set(qn('w:val'),'1'); rpr.append(rtl_el)
        cs_el=OxmlElement('w:cs'); rpr.append(cs_el)
        sz=OxmlElement('w:szCs'); sz.set(qn('w:val'), str(int(size*2))); rpr.append(sz)
    return r
def _emit(doc, blocks, rtl):
    align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    for blk in blocks:
        kind=blk[0]
        if kind=="h1":
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if rtl: _set_rtl_p(p)
            _run(p, blk[1], rtl, bold=True, size=19, color=_BLUE)
        elif kind=="h2":
            p=doc.add_paragraph(); p.alignment=align
            if rtl: _set_rtl_p(p)
            _run(p, blk[1], rtl, bold=True, size=14, color=_TEAL)
        elif kind=="scope":
            p=doc.add_paragraph(); p.alignment=align
            if rtl: _set_rtl_p(p)
            _run(p, blk[1], rtl, italic=True, size=9.5, color=_MUT)
        elif kind=="p":
            p=doc.add_paragraph(); p.alignment=(WD_ALIGN_PARAGRAPH.JUSTIFY if not rtl else WD_ALIGN_PARAGRAPH.RIGHT)
            if rtl: _set_rtl_p(p)
            for text,bold in blk[1]: _run(p, text, rtl, bold=bold, size=11, color=_INK)
        elif kind=="bul":
            for item in blk[1]:
                if not item: continue
                p=doc.add_paragraph(style="List Bullet"); p.alignment=align
                if rtl: _set_rtl_p(p)
                _run(p, item, rtl, size=11, color=_INK)
        elif kind=="table":
            headers, rows = blk[1], blk[2]
            tbl=doc.add_table(rows=1, cols=len(headers)); tbl.style="Light Grid Accent 1"
            if rtl:
                tblPr=tbl._tbl.tblPr; bv=OxmlElement('w:bidiVisual'); tblPr.append(bv)
            for j,h in enumerate(headers):
                c=tbl.rows[0].cells[j]; c.text=""
                pp=c.paragraphs[0]; pp.alignment=align
                if rtl: _set_rtl_p(pp)
                _run(pp, str(h), rtl, bold=True, size=10, color=_TEAL)
            for row in rows:
                cells=tbl.add_row().cells
                for j,val in enumerate(row):
                    cells[j].text=""; pp=cells[j].paragraphs[0]; pp.alignment=align
                    if rtl: _set_rtl_p(pp)
                    _run(pp, str(val), rtl, size=10, color=_INK)
        elif kind=="verse":
            ref, runs = blk[1], blk[2]
            p=doc.add_paragraph(); p.alignment=align
            if rtl: _set_rtl_p(p)
            _run(p, ref+"  —  ", rtl, bold=True, size=10, color=_MUT)
            pv=doc.add_paragraph(); pv.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _set_rtl_p(pv)
            for tok,hit in runs:
                _run(pv, tok+" ", True, bold=hit, size=13, color=(_BLUE if hit else _INK))

def build_docx(D):
    doc=Document()
    if D.get("empty"):
        doc.add_paragraph("No occurrences for this root.")
    else:
        for k,lang in enumerate(["en","fa","ar"]):
            if k>0: doc.add_page_break()
            _emit(doc, _blocks(lang, D), rtl=(lang!="en"))
    bio=io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()
