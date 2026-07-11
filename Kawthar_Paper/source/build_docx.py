# -*- coding: utf-8 -*-
"""Build a .docx from markdown via pandoc, with single line-spacing + centered page-number footer.
Usage: python build_docx.py <input.md> <output.docx> <title> [rtl]"""
import sys, subprocess, os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

md, out, title = sys.argv[1], sys.argv[2], sys.argv[3]
rtl = len(sys.argv)>4 and sys.argv[4]=='rtl'
ref = "/tmp/ref_%s.docx" % ("rtl" if rtl else "ltr")

# ---- start from pandoc's DEFAULT reference (keeps table/figure styles), then modify ----
subprocess.run(['pandoc','--print-default-data-file','reference.docx'], stdout=open(ref,'wb'), check=True)
d = Document(ref)
for sname in ['Normal','Body Text','First Paragraph','Compact','Footer','Header']:
    try:
        st=d.styles[sname]; pf=st.paragraph_format
        pf.line_spacing=1.0; pf.space_after=Pt(6); pf.space_before=Pt(0)
    except KeyError:
        pass
# default font
try:
    n=d.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11)
except KeyError: pass
sec=d.sections[0]
foot=sec.footer; foot.is_linked_to_previous=False
p=foot.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run()
for t,val in [('begin',None),('instr','PAGE'),('end',None)]:
    if t=='instr':
        e=OxmlElement('w:instrText'); e.set(qn('xml:space'),'preserve'); e.text=' PAGE '
    else:
        e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),t)
    r._r.append(e)
if rtl:
    # mark default paragraphs bidi (best-effort; pandoc -V dir=rtl also set)
    pass
d.save(ref)

cmd=['pandoc', md, '-o', out, '--reference-doc', ref, '--toc', '--toc-depth=2']
if rtl: cmd += ['-V','dir=rtl','-V','lang=fa']
else:   cmd += ['-V','lang=en']
r=subprocess.run(cmd, capture_output=True, text=True)
print("pandoc rc", r.returncode, (r.stderr or "")[-400:])
print("exists:", os.path.exists(out), (os.path.getsize(out) if os.path.exists(out) else 0), out)
